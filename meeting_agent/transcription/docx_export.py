"""将转录处理结果导出为 .docx 文件"""

from pathlib import Path
from typing import Any, Dict

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_to_docx(
    result: Dict[str, Any],
    output_path: str | Path,
) -> Path:
    """将结构化会议纪要导出为 .docx 文件。

    Args:
        result: 会议纪要数据字典（符合 MeetingOutput 格式）。
        output_path: 输出文件路径。

    Returns:
        生成的 .docx 文件路径。
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # 标题
    meeting_text = str(result.get("meeting", "")).strip()
    lines = meeting_text.split("\n")

    if lines:
        title = lines[0].strip()
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(lines) > 1:
        date_text = lines[1].strip()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(date_text)
        run.bold = True
        run.font.size = Pt(11)

    # 会议正文
    if len(lines) > 2:
        doc.add_paragraph()  # spacing
        body_text = "\n".join(lines[2:]).strip()
        for paragraph in body_text.split("\n\n"):
            p = doc.add_paragraph(paragraph.strip())
            p.paragraph_format.first_line_indent = Inches(0.28)

    doc.add_paragraph()  # --- separator

    # 会议信息
    doc.add_heading("会议信息", level=2)
    info_table = doc.add_table(rows=3, cols=2, style="Light List Accent 1")
    info_table.cell(0, 0).text = "会议时间"
    info_table.cell(0, 1).text = str(result.get("meeting_date", "未明确"))
    info_table.cell(1, 0).text = "推送部门"
    info_table.cell(1, 1).text = "、".join(result.get("push_dept", [])) or "未明确"
    info_table.cell(2, 0).text = "推送用户"
    info_table.cell(2, 1).text = "、".join(result.get("push_user", [])) or "未明确"

    # 日程安排
    schedules = result.get("schedules", []) or []
    if schedules:
        doc.add_heading("日程安排", level=2)
        schedule_table = doc.add_table(rows=len(schedules) + 1, cols=4, style="Light List Accent 1")
        # Header
        for i, header in enumerate(["标题", "负责人", "时间", "备注"]):
            schedule_table.cell(0, i).text = header
        # Data
        for idx, item in enumerate(schedules):
            schedule_table.cell(idx + 1, 0).text = str(item.get("title", ""))
            owners = item.get("owner", []) or []
            schedule_table.cell(idx + 1, 1).text = "、".join(owners) if owners else ""
            times = f"{item.get('start_time', '')} - {item.get('end_time', '')}"
            schedule_table.cell(idx + 1, 2).text = times
            schedule_table.cell(idx + 1, 3).text = str(item.get("description", ""))

    doc.save(str(output_path))
    return output_path
