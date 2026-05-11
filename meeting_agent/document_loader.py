# 导入路径处理工具，专门管理文件路径（跨Windows/Mac都能用）
from pathlib import Path
# 导入python-docx库，用来读取Word文档（.docx格式）
from docx import Document

# 定义函数：输入文件路径，输出读取到的所有纯文本
# file_path 支持字符串 或 Path 对象
def load_docx_text(file_path: str | Path) -> str:
    # 把传入的路径统一转成 Path 对象（方便后续操作）
    path = Path(file_path)

    # 1. 安全检查：文件不存在就报错
    if not path.exists():
        raise FileNotFoundError(f"文件不存在：{path}")

    # 2. 安全检查：必须是 .docx 文件，不支持旧版 .doc
    if path.suffix.lower() != ".docx":
        raise ValueError("当前版本只支持 .docx 文件。")

    # 打开Word文档
    doc = Document(path)
    # 创建空列表，存放所有读取到的文本
    texts: list[str] = []

    # 3. 读取Word里的【所有段落】（正常文字内容）
    for para in doc.paragraphs:
        # 去掉文字前后的空格、换行
        text = para.text.strip()
        # 只添加非空的有效内容
        if text:
            texts.append(text)

    # 4. 读取Word里的【所有表格】（很多会议纪要会用表格排版）
    for table in doc.tables:
        for row in table.rows:  # 遍历每一行
            row_text = []
            for cell in row.cells:  # 遍历每一格
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            # 把一行的内容用 | 符号连接起来，变成 "内容1 | 内容2 | 内容3"
            if row_text:
                texts.append(" | ".join(row_text))

    # 5. 把所有内容拼接成一整段纯文本，用换行分隔
    content = "\n".join(texts).strip()

    # 6. 如果读到的内容是空的，报错提醒
    if not content:
        raise ValueError("docx 中没有读取到有效文本。")

    # 返回最终读取到的所有文本
    return content